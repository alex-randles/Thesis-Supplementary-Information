import os
from docx import Document
from rdflib import *
from collections import defaultdict
import docx
from numpy import unique


# code which creates a table detailing quality issues detected in experiment 1


MQV = Namespace("http://mappingQualityVocabulary.org/")
MQIO = Namespace("https://w3id.org/MQIO#")

# row_cells[4].text = "Mapping"
# row_cells[5].text = "Report"


mapping_descriptions = {
    "Student Project": ["group-03-r2rml", "group-01-r2rml"],
    "OSi": ["OSi-Mapping1", "OSi-Mapping2"],
    "GeoHive": ["Townland_20m", "LEA_20m"],
    "FAIRVASC": ["FAIRVASC-Mapping3"],
}

def add_prefix(iri):
    for ns_prefix, namespace in g.namespaces():
        if iri.startswith(namespace):
            result = ns_prefix + ":" + iri.split(namespace)[-1]
            return result
    return iri


def create_summary_table():
    table = document.add_table(rows=0, cols=4)
    row_cells = table.add_row().cells
    row_cells[0].text = "#"
    row_cells[1].text = "Source"
    row_cells[2].text = "Total Issues"
    row_cells[3].text = "Description of Issue"
    directory = os.fsencode(".")
    i = 0
    for file in os.listdir(directory):
        mapping_directory = os.fsdecode(file)
        mapping_id = mapping_directory.split("(")[0]
        mapping_project = mapping_directory.split("(")[-1].split(")")[0]
        if os.path.isdir(file):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            project_name = [name for name, mappings in mapping_descriptions.items() if mapping_project in mappings]
            if project_name:
                description = "{} ({})".format(project_name[0], mapping_project)
            else:
                description = mapping_project
            row_cells[1].text = description
            sub_directory = os.fsencode(file)
            str_sub_directory =  sub_directory.decode('utf-8')
            for sub_directory_file in os.listdir(sub_directory):
                str_file_name = sub_directory_file.decode('utf-8')
                validation_report = "./" + str_sub_directory + "/validation_report_3.ttl"
                if str_file_name != "validation_report_3.ttl" and str_file_name != "refined_mapping.ttl" and str_file_name.endswith(".ttl"):
                    mapping_file = os.path.basename(sub_directory_file).decode('utf-8')
                    mapping_path = str_sub_directory + "/" + mapping_file
                    print(validation_report)
                    violation_count = 0
                    violation_descriptions = []
                    if os.path.exists(validation_report):
                        global g
                        g = Graph().parse(validation_report, format="ttl")
                        for s, p, o in g.triples((None, RDF.type, MQV.MappingViolation)):
                            metric = list(g.objects(s, MQV.isDescribedBy))[0]
                            metric_id = metric.split("c")[-1]
                            violation_value = list(g.objects(s, MQV.hasValue))[0]
                            prefixed_value = add_prefix(violation_value)
                            violation_count += 1
                            violation_descriptions.append("{} ({})".format(prefixed_value, metric_id))
                        for s, p, o in g.triples((None, RDF.type, MQIO.MappingViolation)):
                            metric = list(g.objects(s, MQIO.isDescribedBy))[0]
                            metric_id = metric.split("c")[-1]
                            metric_id = metric_id.replace("#", "")
                            object_values = list(g.objects(s, MQIO.hasObjectValue))
                            if object_values:
                                violation_value = object_values[0]
                            object_values = list(g.objects(s, MQIO.hasLiteralValue))
                            if object_values:
                                violation_value = object_values[0]
                            prefixed_value = add_prefix(violation_value)
                            violation_count += 1
                            violation_descriptions.append("{} ({})".format(prefixed_value, metric_id))
                    else:
                        print("Does not exist")
                    if violation_descriptions:
                        row_cells[3].text = "•" + "\n• ".join(violation_descriptions)
                    else:
                        row_cells[3].text = "N/a"
                    row_cells[2].text = str(violation_count)
                    p = document.add_paragraph()
                    # mapping_url = "https://github.com/alex-randles/Evaluation-1/blob/main/" + mapping_path
                    # validation_report_url = "https://github.com/alex-randles/Evaluation-1/blob/main/" + str_sub_directory + "/validation_report.ttl"
                    # row_cells[4].text = mapping_url
                    # row_cells[5].text = validation_report_url
            i += 1

def create_complexity_table():
    table = document.add_table(rows=0, cols=5)
    directory = os.fsencode("./mappings")
    row_cells = table.add_row().cells
    row_cells[0].text = "#"
    row_cells[1].text = "Line Numbers"
    row_cells[2].text = "Triple Maps"
    row_cells[3].text = "Sources"
    row_cells[4].text = "Distinct Ontologies"
    source_counts = defaultdict(dict)
    for file in os.listdir(directory):
        mapping_directory = os.fsdecode(file)
        mapping_id = mapping_directory.split("(")[0]
        mapping_project = mapping_directory.split("(")[-1].split(")")[0]
        if os.path.isdir(file):
            project_name = [name for name, mappings in mapping_descriptions.items() if mapping_project in mappings]
            sub_directory = os.fsencode(file)
            str_sub_directory = sub_directory.decode('utf-8')
            for sub_directory_file in os.listdir(sub_directory):
                str_file_name = sub_directory_file.decode('utf-8')
                validation_report = "./" + str_sub_directory + "/validation_report.ttl"
                if str_file_name.endswith(".ttl") and "foodprint" not in str_file_name:
                    if "refined_mapping" not in str_file_name and "validation_report" not in str_file_name and "mapping_individuals.ttl" not in str_file_name:
                        print(str_sub_directory + "/" + str_file_name)
                        g = Graph().parse(str_sub_directory + "/" + str_file_name, format="ttl")
                        if len(list(g.triples((None, URIRef("http://www.w3.org/ns/r2rml#subjectMap"), None)))) > 1:
                            mapping_file = os.path.basename(sub_directory_file).decode('utf-8')
                            mapping_path = str_sub_directory + "/" + mapping_file
                            row_cells[0].text = mapping_path
                            if "FAIRVASC" in mapping_path:
                                source = "Fairvasc"
                            elif "OSi" in mapping_path:
                                source = "OSi"
                            elif "Beyond" in mapping_path:
                                source = "Beyond2022"
                            elif "group" in mapping_path:
                                source = "Students"
                            else:
                                source = "data.geohive.ie"
                            f = open(mapping_path)
                            line_numbers = len(f.readlines()) - 7
                            print(mapping_path, "whshshzh")
                            if "line_number" not in source_counts[source]:
                                source_counts[source]["line_number"] = line_numbers
                            else:
                                source_counts[source]["line_number"] =  source_counts[source]["line_number"] + line_numbers
                            try:
                                print(mapping_path)
                                print("shsh")
                                g = Graph().parse(mapping_path, format="ttl")
                                triple_maps = len(
                                    list(g.triples((None, URIRef("http://www.w3.org/ns/r2rml#subjectMap"), None))))
                                sources = list(g.objects(None, URIRef("http://www.w3.org/ns/r2rml#logicalTable")))
                                sources.append(list(g.objects(None, URIRef("http://www.w3.org/ns/r2rml#sqlQuery"))))

                                # distinct ontologies
                                classes =  list(g.triples((None, URIRef("http://www.w3.org/ns/r2rml#class"), None)))
                                predicates =  list(g.triples((None, URIRef("http://www.w3.org/ns/r2rml#predicate"), None)))
                                distinct_count = 0
                                distinct_cache = []
                                if "namespaces" not in source_counts[source]:
                                    source_counts[source]["namespaces"] = []
                                if "sources" not in source_counts[source]:
                                    source_counts[source]["sources"] = sources
                                else:
                                    source_counts[source]["sources"].append(sources)
                                for class_value in classes:
                                    class_value = class_value[2]
                                    if "#" in class_value:
                                        namespace = "#".join(class_value.split("#")[:-1])
                                        source_counts[source]["namespaces"].append(namespace)
                                    else:
                                        namespace = "/".join(class_value.split("/")[:-1])
                                        source_counts[source]["namespaces"].append(namespace)
                                for property_value in predicates:
                                    property_value = property_value[2]
                                    if "#" in property_value:
                                        namespace = "/".join(property_value.split("#")[:1])
                                        source_counts[source]["namespaces"].append(namespace)
                                    else:
                                        namespace = "/".join(property_value.split("/")[:-1])
                                        source_counts[source]["namespaces"].append(namespace)
                                values = ["triple_maps"]
                                for label in values:
                                    print(label)
                                    if label not in source_counts[source]:
                                        source_counts[source][label] = locals()[label]
                                    else:
                                        source_counts[source][label] += locals()[label]
                                    print(source_counts)
                            except Exception as e:
                                print(e)
                                pass

    for source, counts in source_counts.items():
        row_cells = table.add_row().cells
        sources = [str(s) for s in counts["sources"]]
        row_cells[0].text = source
        row_cells[1].text = str(counts["line_number"])
        row_cells[2].text = str(counts["triple_maps"])
        row_cells[3].text = str(len(unique(sources)))
        row_cells[4].text = str(len(unique(counts["namespaces"])))

    print(source_counts)
                        # row_cells[4].text = str(line_numbers)

if __name__ == "__main__":
    document = Document()
    create_summary_table()
    document.save("experiment_1_results.docx")
